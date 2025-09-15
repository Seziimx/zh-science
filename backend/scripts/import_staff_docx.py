from __future__ import annotations
import json
from typing import List, Dict, Optional
from pathlib import Path

from docx import Document  # type: ignore

# Import from the backend package
from backend.app.db import SessionLocal
from backend.app.models import User

# Columns we try to detect (kaz/rus variants)
COL_NAME_KEYS = [
    "оқытушының аты-жөні", "аты-жөні", "фио", "фамилия имя отчество",
]
COL_POS_KEYS = [
    "қызметі", "должность",
]
COL_DEGREE_KEYS = [
    "ғылыми атағы", "ғылыми атағы, дәрежесі", "степень", "звание",
]
COL_DEPT_KEYS = [
    "кафедра", "кафедрасы",
]


def _clean(s: Optional[str]) -> str:
    return (s or "").replace("\xa0", " ").strip()


def _detect_col_indexes(headers: List[str]) -> Dict[str, int]:
    idx = {}
    lowered = [h.lower() for h in headers]

    def find(keys: List[str]) -> int:
        for k in keys:
            for i, h in enumerate(lowered):
                if k in h:
                    return i
        return -1

    idx["name"] = find(COL_NAME_KEYS)
    idx["position"] = find(COL_POS_KEYS)
    idx["degree"] = find(COL_DEGREE_KEYS)
    idx["department"] = find(COL_DEPT_KEYS)
    return idx


def _name_variants(full_name: str) -> List[str]:
    parts = [p for p in _clean(full_name).split() if p]
    if not parts:
        return []
    last = parts[0]
    initials = " ".join([f"{p[0]}." for p in parts[1:]])
    cand = set()
    cand.add(full_name)
    if initials:
        cand.add(f"{last} {''.join([p[0]+'.' for p in parts[1:]])}")
        cand.add(f"{last} {initials}")
        cand.add(f"{''.join([p[0]+'.' for p in parts[1:]])} {last}")
    return sorted({c.strip() for c in cand if c.strip()})


def import_staff_from_docx(path: Path) -> Dict[str, int]:
    doc = Document(str(path))
    current_faculty = ""
    created = 0
    updated = 0

    db = SessionLocal()
    try:
        # Iterate through document: paragraphs + tables in sequence is hard with python-docx,
        # so we will assume: a faculty heading paragraph directly precedes its table.
        # Iterate over all tables, and for each table, take nearest previous non-empty paragraph text as faculty.
        para_texts = [p.text.strip() for p in doc.paragraphs]

        # Build mapping from table index to nearest previous non-empty paragraph text
        faculty_by_table: Dict[int, str] = {}
        pi = len(para_texts) - 1
        # We scan once to collect paragraph indices; python-docx doesn't provide a unified run order reliably,
        # so heuristic: for each table, look backwards in document.paragraphs for the nearest non-empty line seen so far.
        # Simpler: just keep a moving pointer while iterating paragraphs and tables is complex -> fallback: use last non-empty paragraph seen before calling this function.
        # We'll instead walk document._body._body like-level content if available is complex; use simpler heuristic below.
        last_heading = ""
        p_iter = iter(doc.element.body)
        tbl_idx = 0
        for el in p_iter:
            tag = el.tag.rsplit('}', 1)[-1]
            if tag == 'p':
                # paragraph
                from docx.text.paragraph import Paragraph  # type: ignore
                paragraph = Paragraph(el, doc)
                txt = paragraph.text.strip()
                if txt:
                    last_heading = txt
            elif tag == 'tbl':
                faculty_by_table[tbl_idx] = last_heading
                tbl_idx += 1

        # Now iterate tables with detected faculty
        for t_idx, table in enumerate(doc.tables):
            faculty = _clean(faculty_by_table.get(t_idx, current_faculty) or current_faculty)
            if faculty:
                current_faculty = faculty

            if not table.rows:
                continue
            # Read header row
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            idx = _detect_col_indexes(headers)
            if idx.get("name", -1) == -1:
                # Not a staff table
                continue

            for row in table.rows[1:]:
                cells = [c.text for c in row.cells]
                name = _clean(cells[idx["name"]]) if idx.get("name", -1) != -1 else ""
                if not name or name.lower() in ("№", "н/д"):
                    continue
                position = _clean(cells[idx["position"]]) if idx.get("position", -1) != -1 else ""
                degree = _clean(cells[idx["degree"]]) if idx.get("degree", -1) != -1 else ""
                department = _clean(cells[idx["department"]]) if idx.get("department", -1) != -1 else ""

                # Upsert user by full_name
                u = db.query(User).filter(User.full_name == name).first()
                if not u:
                    u = User(
                        full_name=name,
                        email=None,
                        role="teacher",
                        faculty=faculty or "",
                        department=department or "",
                        position=position or "",
                        degree=degree or "",
                        login=name.replace(" ", ".").lower(),
                        password_hash="",  # irrelevant here; admin can set later
                        name_variants=json.dumps(_name_variants(name), ensure_ascii=False),
                        active=1,
                    )
                    db.add(u)
                    created += 1
                else:
                    # Update fields
                    changed = False
                    for f, v in {
                        "faculty": faculty or "",
                        "department": department or "",
                        "position": position or "",
                        "degree": degree or "",
                    }.items():
                        if getattr(u, f) != v:
                            setattr(u, f, v)
                            changed = True
                    if changed:
                        updated += 1
                db.flush()
        db.commit()
    finally:
        db.close()

    return {"created": created, "updated": updated}


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python -m backend.scripts.import_staff_docx <path_to_docx>")
        sys.exit(1)
    p = Path(sys.argv[1])
    if not p.exists():
        print(f"File not found: {p}")
        sys.exit(1)
    res = import_staff_from_docx(p)
    print(res)
